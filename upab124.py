#python
import os
import json
import requests
import streamlit as st
from datetime import datetime
from docx import Document

OLLAMA_URL = "http://10.144.177.192:12345/api/generate"
MODEL_NAME = "gpt-oss:20b"
OUTPUT_DIR = "generated_docs"

os.makedirs(OUTPUT_DIR, exist_ok=True)

st.set_page_config(page_title="Procurement Justification Generator", layout="wide")
st.title("AI Procurement Justification Generator")

# ── Scrollable form CSS ──────────────────────────────────────────────────────
st.markdown("""
<style>
div[data-testid="stForm"] {
    max-height: 68vh;
    overflow-y: auto;
    border: 1px solid rgba(250,250,250,0.15);
    border-radius: 10px;
    padding: 1.2rem 1.4rem 1rem 1.4rem;
}
div[data-testid="stForm"]::-webkit-scrollbar {
    width: 5px;
}
div[data-testid="stForm"]::-webkit-scrollbar-track {
    background: transparent;
}
div[data-testid="stForm"]::-webkit-scrollbar-thumb {
    background: rgba(250,250,250,0.2);
    border-radius: 4px;
}
div[data-testid="stForm"]::-webkit-scrollbar-thumb:hover {
    background: rgba(250,250,250,0.4);
}
</style>
""", unsafe_allow_html=True)
# ────────────────────────────────────────────────────────────────────────────

if "form_submitted" not in st.session_state:
    st.session_state.form_submitted = False

if "conversation" not in st.session_state:
    st.session_state.conversation = []

if "answers" not in st.session_state:
    st.session_state.answers = {}

if "current_question" not in st.session_state:
    st.session_state.current_question = None

if "justification_generated" not in st.session_state:
    st.session_state.justification_generated = False

if "final_justification" not in st.session_state:
    st.session_state.final_justification = ""

if "question_count" not in st.session_state:
    st.session_state.question_count = 0


def call_ollama(prompt):
    payload = {
        "model": MODEL_NAME,
        "prompt": prompt,
        "stream": False
    }

    response = requests.post(OLLAMA_URL, json=payload)

    if response.status_code != 200:
        raise Exception(response.text)

    return response.json()["response"]


def build_question_prompt(form_data, conversation, answers):
    return f"""
You are an enterprise procurement assistant.

Your goal is to collect enough information to generate a detailed procurement justification note.

Rules:
1. Ask only ONE simple question at a time.
2. Questions must dynamically depend on previous answers.
3. Ask questions regarding (but not these questions exactly, reframe them and based on the replies):
   - Is this a new requirement and if not then how was it handled earlier?
   - How many number and how did the user arrive at that specific number ((s)he has to provide the calculations)
   - Why were the alternatives not considered (provide those alternatives)?
   - What is the broad technical requirement of this item (ask those technical requirements based on the item under consideration) 
   - Why were these specific technical specs needed and how did the user arrive at these.
   - What performance expectations exist (throughput, latency, scalability)?  
   - Which systems or components must it integrate with?    
   - Where will it be deployed (cloud, on‑prem, edge, field)?  
   - Are there any hardware or software constraints?  
   - What are the maintenance, support, and lifecycle expectations?
4. Do not ask repeated questions.
5. Ask 10 questions and then respond ONLY with:

JUSTIFICATION_COMPLETE

FORM DATA:
{json.dumps(form_data, indent=2)}

ANSWERS:
{json.dumps(answers, indent=2)}

CONVERSATION:
{conversation}
"""


def build_justification_prompt(form_data, answers):
    return f"""
You are an enterprise procurement documentation specialist.

Generate a detailed procurement justification note. Don't include the person name. The justification is based on the 
technical aspects only. 

Requirements:
1. Output only continuous professional paragraphs.
2. No headings.
3. No bullet points.
4. No numbered lists.
5. The output must span multiple detailed paragraphs.
6. Use formal corporate language.
7. Naturally include:
   - technical justification
   - technical calculations
   - alternatives considered
   - existing situation and solutions
   - number of items

FORM DATA:
{json.dumps(form_data, indent=2)}

QUESTION ANSWERS:
{json.dumps(answers, indent=2)}
"""


def generate_docx(employee_name, item_name, justification_text):
    document = Document()

    document.add_heading("Procurement Justification Note", level=1)

    document.add_paragraph(f"Employee: {employee_name}")
    document.add_paragraph(f"Item: {item_name}")
    document.add_paragraph(
        f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
    )

    for para in justification_text.split("\n\n"):
        para = para.strip()

        if para:
            document.add_paragraph(para)

    filename = (
        f"{employee_name.replace(' ', '_')}_"
        f"{item_name.replace(' ', '_')}.docx"
    )

    filepath = os.path.join(OUTPUT_DIR, filename)

    document.save(filepath)

    return filepath


if not st.session_state.form_submitted:

    with st.form("procurement_form"):

        # ── Original fields — untouched ──────────────────────────────────────
        employee_name = st.text_input("Employee Name")

        category = st.selectbox(
            "Category",
            ["IT Hardware", "Software", "Mechanical", "Other"]
        )

        item_name = st.text_input("Item / Service Name")

        estimated_cost = st.number_input(
            "Estimated Cost (in Lakhs)",
            min_value=0.0,
            step=100.0
        )

        urgency = st.selectbox(
            "Urgency",
            ["Low", "Medium", "High", "Critical"]
        )

        procurement_summary = st.text_area("Short Procurement Summary")

        st.divider()

        # ── Section 1 : Initiator Details ────────────────────────────────────
        with st.expander("1. Initiator Details"):

            initiator_name = st.text_input("Initiating Officer Name")

            designation = st.selectbox(
                "Designation",
                [
                    "Scientist B",
                    "Scientist C",
                    "Scientist D",
                    "Scientist E",
                    "Scientist F",
                    "Scientist G"
                ]
            )

            nature_of_item = st.text_input("Nature of Item")

        # ── Section 2 : Budget Details ───────────────────────────────────────
        with st.expander("2. Budget Details"):

            financial_year        = st.text_input("Financial Year")
            fund_1                = st.text_input("Fund 1")
            fund_2                = st.text_input("Fund 2")
            budget_head           = st.text_input("Budget Head")
            budget_core           = st.text_input("Budget Core")
            unit_core             = st.text_input("Unit Core")
            budget_serial_number  = st.text_input("Budget Serial Number")
            forecasted_item       = st.text_input("Forecasted Item")

        # ── Section 3 : Procurement Method Details ───────────────────────────
        with st.expander("3. Procurement Method Details"):

            procurement_method = st.selectbox(
                "Procurement Method",
                ["DD1", "DD2", "DD3", "DD4"]
            )

        # ── Section 4 : Item Details ─────────────────────────────────────────
        with st.expander("4. Item Details"):

            item_serial_number  = st.text_input("Serial Number",     key="item_serial")
            item_code           = st.text_input("Item Code")
            item_type           = st.text_input("Item Type")
            item_nomenclature   = st.text_input("Item Nomenclature")
            item_quantity       = st.text_input("Quantity")
            item_rate           = st.text_input("Rate")
            item_estimated_cost = st.text_input("Estimated Cost")

        # ── Section 5 : Other Inputs (26 fields) ─────────────────────────────
        with st.expander("5. Other Inputs"):

            other_inputs = {}

            for i in range(1, 27):
                other_inputs[f"sentence_{i}"] = st.text_input(
                    f"Sentence {i}",
                    key=f"other_input_{i}"
                )

        # ── Section 6 : Vendors ──────────────────────────────────────────────
        with st.expander("6. Vendors"):

            vendor_serial_number    = st.text_input("Serial Number",           key="vendor_serial")
            vendor_code             = st.text_input("Vendor Code")
            vendor_name             = st.text_input("Vendor Name")
            vendor_address          = st.text_area("Address",                  height=80)
            lab_registration_details= st.text_input("Lab Registration Details")

        # ── Section 7 : RFP Terms ────────────────────────────────────────────
        with st.expander("7. RFP Terms"):

            rfp_term_1 = st.text_input("Indigenous Seller")
            rfp_term_2 = st.text_input("Foreign Seller")
            rfp_term_3 = st.text_input("Advance Payment")

        # ── Section 8 : PAC Accorded ─────────────────────────────────────────
        with st.expander("8. Is PAC Accorded for Item?"):

            pac_accorded = st.selectbox("PAC Accorded", ["Yes", "No"])

        # ── Section 9 : Statement of Case (27 fields) ────────────────────────
        with st.expander("9. Statement of Case"):

            statement_inputs = {}

            for i in range(1, 28):
                statement_inputs[f"field_{i}"] = st.text_input(
                    f"Field {i}",
                    key=f"statement_field_{i}"
                )

        # ── Section 10 : Financial Power ─────────────────────────────────────
        with st.expander("10. Financial Power"):

            financial_power_1 = st.selectbox("Financial Power 1", ["Option 1", "Option 2", "Option 3"])
            financial_power_2 = st.selectbox("Financial Power 2", ["Option 1", "Option 2", "Option 3"])
            financial_power_3 = st.selectbox("Financial Power 3", ["Option 1", "Option 2", "Option 3"])
            financial_power_4 = st.selectbox("Financial Power 4", ["Option 1", "Option 2", "Option 3"])

        # ── Submit ────────────────────────────────────────────────────────────
        submitted = st.form_submit_button("Start AI Interview")

    if submitted:

        form_data = {
            # Original fields
            "employee_name"          : employee_name,
            "category"               : category,
            "item_name"              : item_name,
            "estimated_cost"         : estimated_cost,
            "urgency"                : urgency,
            "procurement_summary"    : procurement_summary,
            # Section 1 — Initiator Details
            "initiator_name"         : initiator_name,
            "designation"            : designation,
            "nature_of_item"         : nature_of_item,
            # Section 2 — Budget Details
            "financial_year"         : financial_year,
            "fund_1"                 : fund_1,
            "fund_2"                 : fund_2,
            "budget_head"            : budget_head,
            "budget_core"            : budget_core,
            "unit_core"              : unit_core,
            "budget_serial_number"   : budget_serial_number,
            "forecasted_item"        : forecasted_item,
            # Section 3 — Procurement Method
            "procurement_method"     : procurement_method,
            # Section 4 — Item Details
            "item_serial_number"     : item_serial_number,
            "item_code"              : item_code,
            "item_type"              : item_type,
            "item_nomenclature"      : item_nomenclature,
            "item_quantity"          : item_quantity,
            "item_rate"              : item_rate,
            "item_estimated_cost"    : item_estimated_cost,
            # Section 5 — Other Inputs
            "other_inputs"           : other_inputs,
            # Section 6 — Vendors
            "vendor_serial_number"   : vendor_serial_number,
            "vendor_code"            : vendor_code,
            "vendor_name"            : vendor_name,
            "vendor_address"         : vendor_address,
            "lab_registration_details": lab_registration_details,
            # Section 7 — RFP Terms
            "rfp_term_1"             : rfp_term_1,
            "rfp_term_2"             : rfp_term_2,
            "rfp_term_3"             : rfp_term_3,
            # Section 8 — PAC Accorded
            "pac_accorded"           : pac_accorded,
            # Section 9 — Statement of Case
            "statement_of_case"      : statement_inputs,
            # Section 10 — Financial Power
            "financial_power_1"      : financial_power_1,
            "financial_power_2"      : financial_power_2,
            "financial_power_3"      : financial_power_3,
            "financial_power_4"      : financial_power_4,
        }

        st.session_state.form_data = form_data
        st.session_state.form_submitted = True

        prompt = build_question_prompt(form_data, [], {})

        first_question = call_ollama(prompt)

        st.session_state.current_question = first_question

        st.rerun()

else:

    for msg in st.session_state.conversation:

        with st.chat_message(msg["role"]):
            st.write(msg["content"])

    if st.session_state.justification_generated:

        st.subheader("Generated Justification")
        st.write(st.session_state.final_justification)

        docx_path = generate_docx(
            st.session_state.form_data["employee_name"],
            st.session_state.form_data["item_name"],
            st.session_state.final_justification
        )

        with open(docx_path, "rb") as file:
            st.download_button(
                label="Download DOCX",
                data=file,
                file_name=os.path.basename(docx_path),
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

    else:

        with st.chat_message("assistant"):
            st.write(st.session_state.current_question)

        user_answer = st.chat_input("Enter your answer")

        if user_answer:

            st.session_state.conversation.append(
                {
                    "role": "assistant",
                    "content": st.session_state.current_question
                }
            )

            st.session_state.conversation.append(
                {
                    "role": "user",
                    "content": user_answer
                }
            )

            question_key = (
                f"question_{st.session_state.question_count + 1}"
            )

            st.session_state.answers[question_key] = {
                "question": st.session_state.current_question,
                "answer": user_answer
            }

            st.session_state.question_count += 1

            next_prompt = build_question_prompt(
                st.session_state.form_data,
                st.session_state.conversation,
                st.session_state.answers
            )

            next_question = call_ollama(next_prompt)

            if "JUSTIFICATION_COMPLETE" in next_question:

                justification_prompt = build_justification_prompt(
                    st.session_state.form_data,
                    st.session_state.answers
                )

                final_justification = call_ollama(
                    justification_prompt
                )

                st.session_state.final_justification = final_justification
                st.session_state.justification_generated = True

            else:
                st.session_state.current_question = next_question

            st.rerun()
