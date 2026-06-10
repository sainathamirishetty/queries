import io
import json
import sqlite3
import hashlib
import requests
import streamlit as st
from datetime import datetime
from docx import Document
import re
import logging

from streamlit_jodit import st_jodit

#login imports
#from ldap3 import Server, Connection, ALL, SIMPLE, SUBTREE
#from ldap3.core.exceptions import LDAPBindError,LDAPException

logger = logging.getLogger(__name__)

def ad_auth(username: str, password: str, client_ip: str):
    """Return (True, display_name) on success, (False, error_msg) on failure."""
    domain = 'hastinapur'
    server_ip = '172.30.1.136'
    user_principal = f"{domain}\\{username}"
    ldap_server = f"ldap://{server_ip}"
    server = Server(ldap_server, get_info=ALL, use_ssl=False)

    try:
        conn = Connection(
            server,
            user=user_principal,
            password=password,
            authentication=SIMPLE,
            auto_bind="NO_TLS"
        )

        if not conn.bind():
            logger.warning(f"AD Auth Failed: {username} from {client_ip}")
            return False, "Invalid credentials"

        # Pull display name (CN) – optional but handy for UI
        search_base = f"DC={domain},DC=res"
        search_filter = f"(&(objectClass=user)(sAMAccountName={username}))"
        conn.search(
            search_base=search_base,
            search_filter=search_filter,
            search_scope=SUBTREE,
            attributes=['displayName']
        )
        if conn.entries:
            dn = conn.entries[0].entry_dn
            cn_match = re.search(r'CN=([^,]+)', dn)
            display_name = cn_match.group(1) if cn_match else username
            logger.info(f"AD Auth Success: {username} ({display_name}) from {client_ip}")
            return True, display_name

        # Fallback – user exists but no CN found
        logger.info(f"AD Auth Success: {username} from {client_ip}")
        return True, username

    except LDAPBindError as e:
        logger.warning(f"AD Auth Failed: {username} from {client_ip} - {str(e)}")
        return False, "Invalid credentials"
    except LDAPException as e:
        logger.warning(f"AD Auth Failed: {username} from {client_ip} - {str(e)}")
        return False, "Invalid credentials"
    except Exception as e:
        logger.warning(f"AD Auth Failed: {username} from {client_ip} - {str(e)}")
        return False, "Invalid credentials"

# ══════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════
OLLAMA_URL = "http://10.144.177.192:12345/api/generate"
MODEL_NAME = "gpt-oss:20b"
DB_PATH    = "procurement.db"

# ── Local testing flag ──────────────────────────────────
# Set USE_LOCAL = True  → bypass AD, use hardcoded credentials below.
# Set USE_LOCAL = False → use real AD authentication.
USE_LOCAL      = True
LOCAL_USERNAME = "abc"
LOCAL_PASSWORD = "123"
# ────────────────────────────────────────────────────────

# ══════════════════════════════════════════════════════════
#  PAGE CONFIG
# ══════════════════════════════════════════════════════════
st.set_page_config(page_title="Procurement Justification Generator", layout="wide")

# ══════════════════════════════════════════════════════════
#  COOKIE MANAGER
# ══════════════════════════════════════════════════════════
try:
    import extra_streamlit_components as stx

    @st.cache_resource
    def _get_cm():
        return stx.CookieManager()

    _cm        = _get_cm()
    COOKIES_OK = True
except Exception:
    _cm        = None
    COOKIES_OK = False


def _set_cookie(key, value):
    if COOKIES_OK:
        _cm.set(key, str(value))


def _get_cookie(key):
    if COOKIES_OK:
        return _cm.get(key)
    return None


def _del_cookie(key):
    if COOKIES_OK:
        try:
            _cm.delete(key)
        except Exception:
            pass


# ══════════════════════════════════════════════════════════
#  DATABASE
# ══════════════════════════════════════════════════════════
def _conn():
    return sqlite3.connect(DB_PATH, check_same_thread=False)


def init_db():
    with _conn() as c:
        # ---- users table (kept for possible future use) ----
        c.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id       INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password TEXT NOT NULL
            )
        """)

        # ---- sessions table -------------------------------------------------
        # If you **do not** need per‑AD‑user isolation, drop the ad_username column.
        c.execute("""
            CREATE TABLE IF NOT EXISTS sessions (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id       INTEGER,                     -- 0 for AD users
                title         TEXT    NOT NULL,
                form_data     TEXT,
                conversation  TEXT,
                answers       TEXT,
                justification TEXT,
                created_at    TEXT DEFAULT (strftime('%Y-%m-%d %H:%M:%S','now','localtime')),
                ad_username   TEXT                         -- optional, can be NULL
            )
        """)
        c.commit()


def db_verify(username, password):
    pw = hashlib.sha256(password.encode()).hexdigest()
    with _conn() as c:
        row = c.execute(
            "SELECT id FROM users WHERE username=? AND password=?",
            (username, pw),
        ).fetchone()
    return row[0] if row else None


def db_get_sessions(user_id):
    with _conn() as c:
        if user_id == 0:   # AD login
            ad_user = st.session_state.username
            return c.execute(
                "SELECT id, title, created_at FROM sessions "
                "WHERE ad_username=? ORDER BY created_at DESC",
                (ad_user,),
            ).fetchall()
        else:
            return c.execute(
                "SELECT id, title, created_at FROM sessions "
                "WHERE user_id=? ORDER BY created_at DESC",
                (user_id,),
            ).fetchall()


def db_save_session(user_id, title, form_data, conversation, answers, justification):
    user_id = user_id or 0
    ad_user = st.session_state.username if user_id == 0 else None
    with _conn() as c:
        cur = c.execute(
            """
            INSERT INTO sessions
            (user_id, title, form_data, conversation, answers,
             justification, ad_username)
            VALUES (?,?,?,?,?,?,?)
            """,
            (
                user_id,
                title,
                json.dumps(form_data),
                json.dumps(conversation),
                json.dumps(answers),
                justification,
                ad_user,
            ),
        )
        c.commit()
        return cur.lastrowid
    
    
def db_rename(session_id, new_title):
    with _conn() as c:
        c.execute(
            "UPDATE sessions SET title=? WHERE id=?", (new_title, session_id)
        )
        c.commit()


def db_delete(session_id):
    with _conn() as c:
        c.execute("DELETE FROM sessions WHERE id=?", (session_id,))
        c.commit()


def db_load(session_id):
    with _conn() as c:
        row = c.execute(
            "SELECT form_data, conversation, answers, justification "
            "FROM sessions WHERE id=?",
            (session_id,),
        ).fetchone()
    if not row:
        return None
    return {
        "form_data":     json.loads(row[0]) if row[0] else {},
        "conversation":  json.loads(row[1]) if row[1] else [],
        "answers":       json.loads(row[2]) if row[2] else {},
        "justification": row[3] or "",
    }


# ══════════════════════════════════════════════════════════
#  SESSION STATE DEFAULTS
# ══════════════════════════════════════════════════════════
_DEFAULTS = dict(
    logged_in=False, user_id=None, username=None,
    form_submitted=False, conversation=[], answers={},
    current_question=None, justification_generated=False,
    final_justification="", question_count=0,
    form_data={}, current_session_id=None,
    unsaved_work=False, sessions_list=[],
    show_save_dialog=False, rename_id=None, rename_title="",
    delete_id=None, pending_new_session=False,
    # ---- form field defaults ----
    f_project_name="",
    f_nature_of_item="",
    f_item_nomenclature="",
    f_technical_parameters="",
    f_any_committee_recommendation="No",
    f_doc_no="",
    f_date=None,
    f_mom_committee_suggestions="",
    f_fresh_purchase="No",
    f_fresh_purchase_purpose_served="",
    f_fresh_purchase_reason="",
    f_previous_supply_order_no="",
    f_previous_supply_order_date=None,
    f_are_items_sensitive="No",
    f_sensitive_items_details="",
    f_sbc_applicable="No",
    f_sbc_doc_no="",
    f_sbc_doc_date=None,
    f_sbc_reason="",
    f_pac_applicable="No",
    f_pac_doc_no="",
    f_pac_doc_date=None,
    f_pac_reason="",
    f_base_of_quantity="",
    f_tender_type="",
    f_tdoc_no="",
    f_tdoc_date=None,
    f_tender_type_reason="",
    f_tender_mode="",
    f_tender_mode_reason="",
    f_bid_type="",
    f_bid_type_reason="",
    f_total_demand_value=0.0,
    f_Gen_justification="",
)
for _k, _v in _DEFAULTS.items():
    if _k not in st.session_state:
        st.session_state[_k] = _v

init_db()


# ══════════════════════════════════════════════════════════
#  FORM FIELD KEYS
#  All widget keys are prefixed with "f_" so they live in
#  session_state and survive reruns without st.form.
# ══════════════════════════════════════════════════════════
# Selectbox option lists (needed by populate_form_fields)
_DESIGNATIONS  = ["Scientist B","Scientist C","Scientist D",
                  "Scientist E","Scientist F","Scientist G"]
_PROC_METHODS  = ["GeM", "Non-GeM"]
_YES_NO        = ["Yes", "No"]
_BID_TYPES     = ["Single Bid", "Two Bid"]
_FIN_POWERS    = ["2.2", "2.7", "2.3", "2.4"]

_SOC_LABELS = [
    ("unit",                  "Unit/Directorate/Office initiating the SOC"),
    ("item_name",             "Name of the item(s)/services being procured"),
    ("justification",         "Justification for procurement"),
    ("authority",             "Authority under which the proposal is being initiated (Financial Power)"),
    ("previous_item",         "Which item was serving the purpose till date"),
    ("current_proposal",      "How the present proposal will serve the purpose"),
    ("broad_purpose",         "Broad purpose of items being procured"),
    ("purchase_category",     "Category of purchase"),
    ("fresh_purchase_purpose","Fresh purchase – how was the purpose being served till date"),
    ("fresh_purchase_reason", "Fresh purchase – why it cannot be served with upgradation of existing items"),
    ("quantity_basis",        "Basis for working out the quantity against each item"),
    ("holding_details",       "Details like authorized holdings, existing holding"),
    ("calc_sheets",           "Calculation sheets/PR documents"),
    ("distribution",          "Proposed distribution of items"),
    ("last_purchase_date",    "Last purchase date"),
    ("budget_quotes",         "Budgetary quotes"),
    ("market_intelligence",   "Market intelligence"),
    ("rates_other_orgs",      "Rates obtained from other organisations"),
    ("professional_eval",     "Professional officers evaluation"),
    ("other_method",          "Any other method"),
    ("is_new_item",           "Is new item"),
    ("major_head",            "Major head, sub head and detailed head"),
    ("fund_certificate",      "Fund availability certificate"),
    ("paying_agency",         "Name of paying agency"),
    ("tender_mode",           "Mode of tendering"),
    ("tender_justification",  "Justification for tendering mode other than OBM"),
    ("draft_rfp",             "Draft RFP – original sentences"),
]


def collect_form_data():
    ss = st.session_state
    return {
        #0.directorate name
        "project_name": ss.get("f_project_name", ""),
        # 1. Nature of Item
        "nature_of_item": ss.get("f_nature_of_item", ""),
        # 2. Item Nomenclature
        "item_nomenclature": ss.get("f_item_nomenclature", ""),
        # 3. Technical Parameters
        "technical_parameters": ss.get("f_technical_parameters", ""),
        # 4. Committee recommendation
        "any_committee_recommendation": ss.get("f_any_committee_recommendation", "No"),
        "doc_no": ss.get("f_doc_no", ""),
        "date": ss.get("f_date").strftime("%Y-%m-%d") if ss.get("f_date") else None,  # Convert to string
        "mom_committee_suggestions": ss.get("f_mom_committee_suggestions", ""),
        # 5. Fresh purchase
        "fresh_purchase": ss.get("f_fresh_purchase", "No"),
        "fresh_purchase_purpose_served": ss.get("f_fresh_purchase_purpose_served", ""),
        "fresh_purchase_reason": ss.get("f_fresh_purchase_reason", ""),
        "previous_supply_order_no": ss.get("f_previous_supply_order_no", ""),
        "previous_supply_order_date": ss.get("f_previous_supply_order_date").strftime("%Y-%m-%d") if ss.get("f_previous_supply_order_date") else None,  # Convert to string
        # 6. Sensitive items
        "are_items_sensitive": ss.get("f_are_items_sensitive", "No"),
        "sensitive_items_details": ss.get("f_sensitive_items_details", ""),
        # 7. SBC
        "sbc_applicable": ss.get("f_sbc_applicable", "No"),
        "sbc_doc_no": ss.get("f_sbc_doc_no", ""),
        "sbc_doc_date": ss.get("f_sbc_doc_date").strftime("%Y-%m-%d") if ss.get("f_sbc_doc_date") else None,  # Convert to string
        "sbc_reason": ss.get("f_sbc_reason", ""),
        #8 PAC
        "pac_applicable": ss.get("f_pac_applicable", "No"),
        "pac_doc_no": ss.get("f_pac_doc_no", ""),
        "pac_reason": ss.get("f_pac_reason", ""),
        "pac_doc_date": ss.get("f_pac_doc_date").strftime("%Y-%m-%d") if ss.get("f_pac_doc_date") else None,
        # 9. Quantity
        #"item_quantity": ss.get("f_item_quantity", 0),
        # 10. Base of quantity
        "base_of_quantity": ss.get("f_base_of_quantity", ""),
        # 11. Proposed distribution
        "proposed_distribution": ss.get("f_proposed_distribution", ""),
        # 12. Tender type
        "tender_type": ss.get("f_tender_type", ""),
        # FIX #2 — tdoc_no and tdoc_date now collected
        "tdoc_no": ss.get("f_tdoc_no", ""),
        "tdoc_date": ss.get("f_tdoc_date").strftime("%Y-%m-%d")
        if ss.get("f_tdoc_date") else None,
        "tender_type_reason": ss.get("f_tender_type_reason", ""),
        # 13. Tender mode
        "tender_mode": ss.get("f_tender_mode", ""),
        "tender_mode_reason": ss.get("f_tender_mode_reason", ""),
        # 14. Bid type
        "bid_type": ss.get("f_bid_type", ""),
        "bid_type_reason": ss.get("f_bid_type_reason", ""),
        # 15. Total demand value
        "total_demand_value": ss.get("f_total_demand_value", 0.0),
        #16 Addl Gen Justification
        "Addl. General Justification details": ss.get("f_Gen_justification")

    }


def populate_form_fields(form_data):
    ss = st.session_state
    # 0.directorate name
    ss["f_project_name"] = form_data.get("project_name", "")
    # 1. Nature of Item
    ss["f_nature_of_item"] = form_data.get("nature_of_item", "")

    # 2. Item Nomenclature
    ss["f_item_nomenclature"] = form_data.get("item_nomenclature", "")

    # 3. Technical Parameters
    ss["f_technical_parameters"] = form_data.get("technical_parameters", "")

    # 4. Committee recommendation
    ss["f_any_committee_recommendation"] = form_data.get("any_committee_recommendation", "No")
    ss["f_doc_no"] = form_data.get("doc_no", "")
    # Handle date conversion - if it's a string, convert back to datetime
    date_val = form_data.get("date")
    ss["f_date"] = datetime.strptime(date_val, "%Y-%m-%d") if isinstance(date_val, str) and date_val else None
    ss["f_mom_committee_suggestions"] = form_data.get("mom_committee_suggestions", "")


    # 5. Fresh purchase
    ss["f_fresh_purchase"] = form_data.get("fresh_purchase", "No")
    ss["f_fresh_purchase_purpose_served"] = form_data.get("fresh_purchase_purpose_served", "")
    ss["f_fresh_purchase_reason"] = form_data.get("fresh_purchase_reason", "")
    ss["f_previous_supply_order_no"] = form_data.get("previous_supply_order_no", "")
    # Handle previous supply order date conversion
    prev_date_val = form_data.get("previous_supply_order_date")
    ss["f_previous_supply_order_date"] = datetime.strptime(prev_date_val, "%Y-%m-%d") if isinstance(prev_date_val, str) and prev_date_val else None

    # 6. Sensitive items
    ss["f_are_items_sensitive"] = form_data.get("are_items_sensitive", "No")
    ss["f_sensitive_items_details"] = form_data.get("sensitive_items_details", "")

    # 7. SBC
    ss["f_sbc_applicable"] = form_data.get("sbc_applicable", "No")
    ss["f_sbc_doc_no"] = form_data.get("sbc_doc_no", "")
    # Handle SBC doc date conversion
    sbc_date_val = form_data.get("sbc_doc_date")
    ss["f_sbc_doc_date"] = datetime.strptime(sbc_date_val, "%Y-%m-%d") if isinstance(sbc_date_val, str) and sbc_date_val else None
    ss["f_sbc_reason"] = form_data.get("sbc_reason", "")
    
    #8.PAC
    ss["f_pac_applicable"] = form_data.get("pac_applicable", "No")
    ss["f_pac_doc_no"] = form_data.get("pac_doc_no", "")
    ss["f_pac_reason"] = form_data.get("pac_reason", "")
    pac_date = form_data.get("pac_doc_date")
    ss["f_pac_doc_date"] = datetime.strptime(pac_date, "%Y-%m-%d") if pac_date else None

    # 9. Quantity
    #ss["f_item_quantity"] = form_data.get("item_quantity", 0)

    # 10. Base of quantity
    ss["f_base_of_quantity"] = form_data.get("base_of_quantity", "")

    # 11. Proposed distribution
    ss["f_proposed_distribution"] = form_data.get("proposed_distribution", "")

    # 12. Tender type
    ss["f_tender_type"] = form_data.get("tender_type", "")
    ss["f_tdoc_no"] = form_data.get("tdoc_no", "")
    tdoc_date_val = form_data.get("tdoc_date")
    ss["f_tdoc_date"] = (
        datetime.strptime(tdoc_date_val, "%Y-%m-%d") if tdoc_date_val else None
    )
    ss["f_tender_type_reason"] = form_data.get("tender_type_reason", "")

    # 13. Tender mode
    ss["f_tender_mode"] = form_data.get("tender_mode", "")
    ss["f_tender_mode_reason"] = form_data.get("tender_mode_reason", "")

    # 14. Bid type
    ss["f_bid_type"] = form_data.get("bid_type", "")
    ss["f_bid_type_reason"] = form_data.get("bid_type_reason", "")

    # 15. Total demand value
    ss["f_total_demand_value"] = form_data.get("total_demand_value", 0.0)
    
    #16. Addl. General Justification
    ss["f_Gen_justification"] = form_data.get("Addl. General Justification details", "")
    
    
# ══════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════
def refresh_sessions():
    """
    Load the list of sessions for the current user.
    AD users (user_id is None) use the sentinel 0.
    """
    uid = st.session_state.user_id or 0
    st.session_state.sessions_list = db_get_sessions(uid)


def reset_work():
    """Clear active-work state and wipe all form field keys."""
    st.session_state.update(
        form_submitted=False, conversation=[], answers={},
        current_question=None, justification_generated=False,
        final_justification="", question_count=0,
        form_data={}, current_session_id=None, unsaved_work=False,
    )
    populate_form_fields({})   # blank out all fields


def group_by_date(sessions):
    grouped = {}
    for sid, title, created_at in sessions:
        dt  = datetime.strptime(created_at, "%Y-%m-%d %H:%M:%S")
        key = dt.strftime("%d %b %Y")
        grouped.setdefault(key, []).append((sid, title, dt.strftime("%H:%M")))
    return grouped


def has_active_work():
    """True if user has filled any field OR started a conversation."""
    any_field = any(
        bool(st.session_state.get(k, ""))
        for k in ("f_initiator_name", "f_item_nomenclature", "f_nature_of_item")
    )
    return any_field or bool(st.session_state.form_data) or bool(st.session_state.conversation)


# ══════════════════════════════════════════════════════════
#  OLLAMA
# ══════════════════════════════════════════════════════════
def call_ollama(prompt):
    r = requests.post(
        OLLAMA_URL,
        json={"model": MODEL_NAME, "prompt": prompt, "stream": False},
        timeout=120,
    )
    if r.status_code != 200:
        raise Exception(r.text)
    return r.json()["response"]


# ══════════════════════════════════════════════════════════
#  PROMPTS
# ══════════════════════════════════════════════════════════
def build_question_prompt(form_data, conversation, answers):
    return f"""
    You are an enterprise procurement assistant and you have technical knowledge on the given product too.
    Your goal is to collect enough information regarding the product to generate a detailed procurement justification note.

Rules:
1. Ask only ONE simple question at a time.
2. Questions must dynamically depend on previous answers.
3. Ask questions only regarding (reframe naturally based on context):
    - Any additional information you want to provide?
    - Why were these specific technical specs needed and how did the user arrive at these.
4. Do not ask duplicate or repeated questions. Only ONE question (not your thinking).
5. Ask at most 3 questions then respond ONLY with: JUSTIFICATION_COMPLETE

    FORM DATA:
    {json.dumps(form_data, indent=2)}

    ANSWERS:
    {json.dumps(answers, indent=2)}

    CONVERSATION:
    {conversation}
    """

def build_justification_prompt(form_data, answers):

    style_block = """
    Write the justification in a **neutral, objective tone**.
    - Use third‑person only; avoid “we”, “I”, or any first‑person references.
    - Do not use promotional adjectives (e.g., “excellent”, “state‑of‑the‑art”).
    - Do not add commentary that is not present in the supplied data.
    - Follow the three‑paragraph structure required by the procurement template.

    Do:
    - State facts exactly as given.
    - Use passive voice only when it improves clarity.
    - Write numbers in digits (e.g., 66.6 kVA).

    Don’t:
    - Insert opinions, recommendations, or justifications not in the data.
    - Use emotive words (e.g., “critical”, “urgent”, “important”).
    - Add headings, bullet points, or extra line breaks.
    """

    # Optional one‑shot example (kept generic)

    prompt = f"""
    {style_block}
    You are an enterprise procurement assistant specialized in drafting official
    justification notes for government and public‑sector purchases.

    Using **only** the information provided in the sections **FORM DATA** and
    **QUESTION ANSWERS**, produce a **single, continuous justification note** that
    consists of **exactly three paragraphs**:

    1. **Technical Justification** – Explain how the need for the product(s) arose,
       referencing relevant project details, operational requirements, or regulatory
       drivers found in the form data and answers.

    2. **Quantity & Distribution Justification** – Detail the required quantity,
       any proposed allocation across departments or locations, and why this amount
       is appropriate given the stated need.

    3. **Tender & Bid Explanation** – The default procurement settings are:
       - Tender type: **GeM**
       - Tender mode: **Open**
       - Bid type: **Two**

       Include this paragraph **only if** any of these defaults differ from the
       values indicated in the input data, and provide the reason supplied by the
       officer.

    **Formatting Rules**
    - Produce three plain paragraphs with no headings, bullet points, or numbered
      lists.
    - Use simple sentences.
    - Use formal, corporate language appropriate for official documentation.
    - Keep the text continuous; do not insert line breaks within a paragraph
      except for the paragraph separation.
    - Don't refer to the form in the text.
    - Generate tables wherever necessary.
    - **ELABORATE THE TEXT**

    **Input Sections**

    FORM DATA:
    {json.dumps(form_data, indent=2)}

    QUESTION ANSWERS:
    {json.dumps(answers, indent=2)}
    """
    return prompt


# ══════════════════════════════════════════════════════════
#  DOCX  (fixed — uses params, returns in-memory buffer)
# ══════════════════════════════════════════════════════════
def generate_docx(item_name, justification_text):
    doc = Document()
    doc.add_heading("Procurement Justification Note", level=1)
    
    doc.add_paragraph(f"Item: {item_name}")
    doc.add_paragraph(f"Generated On: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    for para in justification_text.split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ══════════════════════════════════════════════════════════
#  DIALOGS
# ══════════════════════════════════════════════════════════
@st.dialog(" Save Session")
def save_dialog():
    title = st.text_input(
        "Enter a title for this session",
        placeholder="",
    )
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Save", use_container_width=True):
            if title.strip():
                # Always collect freshest form state before saving
                current_fd = collect_form_data()
                st.session_state.form_data = current_fd
                sid = db_save_session(
                    st.session_state.user_id,
                    title.strip(),
                    current_fd,
                    st.session_state.conversation,
                    st.session_state.answers,
                    st.session_state.final_justification,
                )
                st.session_state.current_session_id = sid
                st.session_state.unsaved_work       = False
                st.session_state.show_save_dialog   = False
                refresh_sessions()
                st.rerun()
            else:
                st.warning("Please enter a title.")
    with c2:
        if st.button("Cancel", use_container_width=True):
            st.session_state.show_save_dialog = False
            st.rerun()


@st.dialog(" Rename Session")
def rename_dialog():
    new = st.text_input("Session title", value=st.session_state.rename_title)
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Update", use_container_width=True):
            if new.strip():
                db_rename(st.session_state.rename_id, new.strip())
                st.session_state.rename_id    = None
                st.session_state.rename_title = ""
                refresh_sessions()
                st.rerun()
            else:
                st.warning("Title cannot be empty.")
    with c2:
        if st.button("Cancel", use_container_width=True):
            st.session_state.rename_id = None
            st.rerun()


@st.dialog(" Delete Session")
def delete_dialog():
    st.warning("Are you sure you want to delete this session? This cannot be undone.")
    c1, c2 = st.columns(2)
    with c1:
        if st.button("Yes, Delete", use_container_width=True):
            sid = st.session_state.delete_id
            if st.session_state.current_session_id == sid:
                reset_work()
            db_delete(sid)
            st.session_state.delete_id = None
            refresh_sessions()
            st.rerun()
    with c2:
        if st.button("Cancel", use_container_width=True):
            st.session_state.delete_id = None
            st.rerun()


@st.dialog(" ")
def unsaved_dialog():
    st.warning("You have unsaved work. What would you like to do?")
    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button(" Save", use_container_width=True):
            st.session_state.pending_new_session = False
            st.session_state.show_save_dialog    = True
            st.rerun()
    with c2:
        if st.button("Continue to create New session", use_container_width=True):
            st.session_state.pending_new_session = False
            reset_work()
            st.rerun()
    with c3:
        if st.button("Cancel", use_container_width=True):
            st.session_state.pending_new_session = False
            st.rerun()


# ══════════════════════════════════════════════════════════
#  SIDEBAR
# ══════════════════════════════════════════════════════════
def show_sidebar():
    with st.sidebar:

        st.markdown(f"###  Welcome, **{st.session_state.username}**")
        st.divider()

        # New Session | Logout
        c1, c2 = st.columns(2)
        with c1:
            if st.button("➕ New Session", use_container_width=True):
                if has_active_work():
                    st.session_state.pending_new_session = True
                else:
                    reset_work()
                st.rerun()
        with c2:
            if st.button(" Logout", use_container_width=True):
                _del_cookie("username")
                _del_cookie("user_id")
                for k in list(st.session_state.keys()):
                    del st.session_state[k]
                st.rerun()

        # Save — always visible when any work exists
        if has_active_work():
            if st.button(" Save", use_container_width=True):
                st.session_state.show_save_dialog = True
                st.rerun()

        st.divider()

        # Chat history
        if not st.session_state.sessions_list:
            st.caption("No saved sessions yet.")
        else:
            grouped = group_by_date(st.session_state.sessions_list)
            for date_str, items in grouped.items():
                st.markdown(f"**{date_str}**")
                for sid, title, time_str in items:
                    is_active     = (sid == st.session_state.current_session_id)
                    col_t, col_m  = st.columns([5, 1])
                    with col_t:
                        if st.button(
                            f"{time_str} - {title}",
                            key=f"sess_{sid}",
                            use_container_width=True,
                            type="primary" if is_active else "secondary",
                        ):
                            data = db_load(sid)
                            if data:
                                populate_form_fields(data["form_data"])
                                st.session_state.form_data               = data["form_data"]
                                st.session_state.conversation            = data["conversation"]
                                st.session_state.answers                 = data["answers"]
                                st.session_state.final_justification     = data["justification"]
                                st.session_state.justification_generated = bool(data["justification"])
                                st.session_state.form_submitted          = (
                                    bool(data["conversation"]) or bool(data["justification"])
                                )
                                st.session_state.current_session_id = sid
                                st.session_state.unsaved_work        = False
                            st.rerun()
                    with col_m:
                        with st.popover("⋮"):
                            if st.button(" Rename", key=f"ren_{sid}", use_container_width=True):
                                st.session_state.rename_id    = sid
                                st.session_state.rename_title = title
                                st.rerun()
                            if st.button(" Delete", key=f"del_{sid}", use_container_width=True):
                                st.session_state.delete_id = sid
                                st.rerun()


# ══════════════════════════════════════════════════════════
#  LOGIN PAGE
# ══════════════════════════════════════════════════════════
def show_login():
    st.markdown(
        "<h1 style='text-align: center;'>Login with RCNET/PIS Details</h1>",
        unsafe_allow_html=True
    )
    _, col, _ = st.columns([1, 1, 1])
    with col:
        st.subheader("Login")
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")

        if st.button("Login", use_container_width=True):
            if USE_LOCAL:
                # ── Local test mode ──
                if username == LOCAL_USERNAME and password == LOCAL_PASSWORD:
                    st.session_state.logged_in = True
                    st.session_state.username  = LOCAL_USERNAME
                    st.session_state.user_id   = None
                    _set_cookie("username", LOCAL_USERNAME)
                    _set_cookie("user_id", "local")
                    refresh_sessions()
                    st.rerun()
                else:
                    st.error("Invalid credentials")
            else:
                # ── AD authentication ──
                client_ip = st.session_state.get("client_ip") or "unknown"
                success, info = ad_auth(username, password, client_ip)

                if success:
                    st.session_state.logged_in = True
                    st.session_state.username  = info
                    st.session_state.user_id   = None
                    _set_cookie("username", info)
                    _set_cookie("user_id", "ad")
                    refresh_sessions()
                    st.rerun()
                else:
                    st.error(info)


# -------------------------------------------------------------
#  MAIN APP
# -------------------------------------------------------------
def show_main_app():
    st.title("AI BASED PROCUREMENT JUSTIFICATION GENERATOR")

    # ── PHASE 1 : NEW FORM (single page, no expanders) ──
    if not st.session_state.form_submitted:
        # 0.  Project/Directorate Name – top of the form
        st.text_input(
            "Project/Directorate Name",
            key="f_project_name",
        )
        # 1. Nature of Item (dropdown)
        col1, col2 = st.columns([1,2])
        with col1:
            st.selectbox(
                "Nature of Item",
                [
                    "ACCESSORIES",
                    "AMMUNITION",
                    "ARBORICULTURE (special services)",
                    "BUILDUP ITEMS",
                    "CAPITAL",
                    "CONSERVANCY (SPECIAL SERVICES)",
                    "CONSUMABLES",
                    "COTS",
                    "DEVELOPMENT",
                    "EQUIPMENTS",
                    "EXPLOSIVE",
                    "FABRICATION",
                    "FIREFIGHTING (SPECIAL SERVICES)",
                    "HARDWARE",
                    "HYGIENE AND MAINTEINANCE (special services)",
                    "IT HARDWARE",
                    "IT SOFTWARE",
                    "MISCELLANEOUS",
                    "RAW MATERIAL",
                    "REPLACEMENT (Repair/Maintenance Contracts)",
                    "SECURITY SERVICES DGR (Special services)",
                    "SERVICES",
                    "SPACES",
                    "SUPPORT SERVICES TO DSC (special services)",
                    "TRANSPORTATION",
                    "WET CANTEEN SERVICES (special services)",
                ],
                key="f_nature_of_item",
            )

        with col2:
        # 2. Item Nomenclature (input box)
            st.text_input("Item Nomenclature", key="f_item_nomenclature")

        # 3. Technical Parameters Calculation (rich text)
        st.markdown("**Technical Parameters Calculation**")
        _tech_val = st_jodit(
            value=st.session_state.get("f_technical_parameters", ""),
            key="jodit_technical_parameters",
            config={
                "height": 300,
                "language": "en",
                "removeButtons": ["mic", "print", "image", "video", "file"],
            },
        )
        if _tech_val is not None:
            st.session_state["f_technical_parameters"] = _tech_val

        # 4. Any Committee Recommendation (radio + conditional fields)
        st.radio(
            "Any Committee Recommendation?",
            ["Yes", "No"],
            key="f_any_committee_recommendation",
        )
        #if st.session_state.f_any_committee_recommendation == "Yes":
        disabled = st.session_state.f_any_committee_recommendation == "No"
        
        col_docno, col_docdate = st.columns([2, 1])
        with col_docno:
            st.text_input("Document No", key="f_doc_no",disabled=disabled,)
        with col_docdate:
            st.date_input("Document Date", key="f_date",disabled=disabled,)
                
        st.text_area(
            "MOM Committee Suggestions",
            height=120,
            key="f_mom_committee_suggestions",
            disabled=disabled,
        )
        

        # 5. Fresh Purchase (radio + conditional fields)
        st.radio(
            "Fresh Purchase?",
            ["Yes", "No"],
            key="f_fresh_purchase",
        )
        disable_prev_supply = st.session_state.f_fresh_purchase == "Yes"
        
        if st.session_state.f_fresh_purchase == "Yes":
            st.text_area(
                "How was the purpose served till now?",
                height=80,
                key="f_fresh_purchase_purpose_served",
            )
        else:
            st.text_area(
                "Why upgradation of existing is not enough?",
                height=80,
                key="f_fresh_purchase_reason",
            )
        col_order_no, col_order_date = st.columns([2, 1])
        
        with col_order_no:
            st.text_input(
                "Previous Supply Order No",
                key="f_previous_supply_order_no",disabled=disable_prev_supply,
            )

        with col_order_date:
            st.date_input(
                "Previous Supply Order Date",
                key="f_previous_supply_order_date",disabled=disable_prev_supply,
            )

        # 6. Are Items Sensitive in Nature (radio + conditional field)
        st.radio(
            "Are Items Sensitive in Nature?",
            ["Yes", "No"],
            key="f_are_items_sensitive",index=0,
        )
        disabled = st.session_state.f_are_items_sensitive == "No"
        st.text_area(
            "Explain how and if FIM will be issued?",
            height=80,
            key="f_sensitive_items_details",
            disabled=disabled,
        )
        

        # 7. SBC (radio + conditional fields + reason)
        st.radio(
            "Is SBC (Specific Brand Ceritificate) applied?",
            ["Yes", "No"],
            key="f_sbc_applicable",
        )
        disabled = st.session_state.f_sbc_applicable == "No"
        
        col_docno, col_docdate = st.columns([2, 1])
            
        with col_docno:
            st.text_input("SBC Doc No", key="f_sbc_doc_no",disabled=disabled,)
            st.text_input("SBC Reason", key="f_sbc_reason",disabled=disabled,)
        
        with col_docdate:
            st.date_input("SBC Doc Date", key="f_sbc_doc_date",disabled=disabled,)
            
            
            
        # 15  PAC – bottom of the form (after Total Demand Value)
        st.radio(
            "Is PAC(Proprietory Article Certificate) Applicable?",
            ["Yes", "No"],
            key="f_pac_applicable",
        )
        pac_disabled = st.session_state.f_pac_applicable == "No"

        col_pac_docno, col_pac_docdate = st.columns([2, 1])
        with col_pac_docno:
            st.text_input(
                "PAC Doc No",
                key="f_pac_doc_no",
                disabled=pac_disabled,
            )
            st.text_input(
                "PAC Reason",
                key="f_pac_reason",
                disabled=pac_disabled,
            )
        with col_pac_docdate:
            st.date_input(
                "PAC Doc Date",
                key="f_pac_doc_date",
                disabled=pac_disabled,
            )
               
            
      
        

        # 8. Quantity (number input)
        #st.number_input(
            #"Quantity",
            #min_value=0,
            #step=1,
            #key="f_item_quantity",
        #)

        # 9. Base of Quantity (text input)
        st.text_area("Line item Quantity & Justification",height=120, key="f_base_of_quantity")

        # 10. Proposed Distribution of Items (long text)
        st.text_area(
            "Proposed Distribution of Items",
            height=120,
            key="f_proposed_distribution",
        )

        #11. tendertype# Clear hidden fields when the user switches back to GeM
        if st.session_state.get("f_tender_type") == "GeM":
            st.session_state.f_tdoc_no = ""
            st.session_state.f_tdoc_date = None
            st.session_state.f_tender_type_reason = ""

        # Now create the widgets
        st.selectbox(
            "Tender Type",
            [
                "GeM",
                "Annual Maint. Cont",
                "Cash Purchase Scientific Equipment",
                "CARS for Research Services",
                "Cash Purchase",
                "Development Contract",
                "Fabrication Order",
                "Imports",
                "CAPSI",
                "Rate Contracts",
                "Supply Order",
            ],
            key="f_tender_type",
        )

        # Conditional block – only when the type is NOT GeM
        disabled = st.session_state.f_tender_type == "GeM"
        
        col_docno, col_docdate = st.columns([2, 1])

        with col_docno:
            st.text_input("Document No", key="f_tdoc_no",disabled=disabled,)
            
            st.text_area(
                "Tender Type Reason",
                height=120,
                key="f_tender_type_reason",disabled=disabled,
            )

        with col_docdate:
            st.date_input("Document Date", key="f_tdoc_date",disabled=disabled,)

            

        # 12. Tender Mode (radio + reason)
        st.radio(
            "Tender Mode",
            ["Single", "Limited", "Open"],
            key="f_tender_mode",
        )
        
        disabled = st.session_state.f_tender_mode == "Open"

        # 3️⃣  Show/hide the reason field based on the selected mode
        st.text_input(
            "Tender Mode Reason",
            key="f_tender_mode_reason",
            disabled=disabled,          
        )
        

        # 13. Bid Type (radio + reason)
        st.radio(
            "Bid Type",
            ["Single", "Two"],
            key="f_bid_type",
        )
        disabled = st.session_state.f_bid_type == "Two"
        
        st.text_input(
            "Bid Type Reason",
            key="f_bid_type_reason",
            disabled=disabled,
        )

        

        # 14. Total Demand Value (number input)
        st.number_input(
            "Total Demand Value (Rupees)",
            min_value=0.0,
            step=1.0,
            key="f_total_demand_value",
        )
        
        # 16. Proposed Distribution of Items (long text)
        st.text_area(
            "Addl. General Justification details",
            height=120,
            key="f_Gen_justification",
        )        
        
        # ────────────────────────────────────────────────────────────────────────
        # 1️⃣  Validation
        # NOT mandatory: f_technical_parameters, f_doc_no, f_date,
        #                f_proposed_distribution, f_Gen_justification
        # ────────────────────────────────────────────────────────────────────────
        missing = []

        ss = st.session_state

        # ── Always mandatory (no conditions) ──
        always_mandatory = [
            "f_project_name",
            "f_nature_of_item",
            "f_item_nomenclature",
            "f_any_committee_recommendation",
            "f_fresh_purchase",
            "f_are_items_sensitive",
            "f_sbc_applicable",
            "f_pac_applicable",
            "f_base_of_quantity",
            "f_tender_type",
            "f_tender_mode",
            "f_bid_type",
            "f_total_demand_value",
        ]
        for key in always_mandatory:
            if not ss.get(key):
                missing.append(key)

        # ── Committee Recommendation ──
        if ss.f_any_committee_recommendation == "Yes":
            if not ss.get("f_mom_committee_suggestions"):
                missing.append("f_mom_committee_suggestions")

        # ── Fresh Purchase conditional fields ──
        if ss.f_fresh_purchase == "Yes":
            # purpose_served is required; previous supply order fields not needed
            if not ss.get("f_fresh_purchase_purpose_served"):
                missing.append("f_fresh_purchase_purpose_served")
        else:
            # reason for not upgrading is required; previous supply order fields required
            if not ss.get("f_fresh_purchase_reason"):
                missing.append("f_fresh_purchase_reason")
            if not ss.get("f_previous_supply_order_no"):
                missing.append("f_previous_supply_order_no")
            if not ss.get("f_previous_supply_order_date"):
                missing.append("f_previous_supply_order_date")

        # ── Sensitive items ──
        if ss.f_are_items_sensitive == "Yes":
            if not ss.get("f_sensitive_items_details"):
                missing.append("f_sensitive_items_details")

        # ── SBC ──
        if ss.f_sbc_applicable == "Yes":
            for key in ("f_sbc_doc_no", "f_sbc_doc_date", "f_sbc_reason"):
                if not ss.get(key):
                    missing.append(key)

        # ── PAC ──
        if ss.f_pac_applicable == "Yes":
            for key in ("f_pac_doc_no", "f_pac_doc_date", "f_pac_reason"):
                if not ss.get(key):
                    missing.append(key)

        # ── Tender Type: Non-GeM requires extra fields ──
        if ss.f_tender_type != "GeM":
            for key in ("f_tdoc_no", "f_tdoc_date", "f_tender_type_reason"):
                if not ss.get(key):
                    missing.append(key)

        # ── Tender Mode: reason required unless Open ──
        if ss.f_tender_mode != "Open":
            if not ss.get("f_tender_mode_reason"):
                missing.append("f_tender_mode_reason")

        # ── Bid Type: reason required unless Two ──
        if ss.f_bid_type != "Two":
            if not ss.get("f_bid_type_reason"):
                missing.append("f_bid_type_reason")

        # 2️⃣  Show warning or the Start button
        if missing:
            # Show human-readable field names (strip f_ prefix, replace _ with space)
            readable = [k[2:].replace("_", " ").title() for k in missing]
            st.warning(f"Please fill in the mandatory fields: {', '.join(readable)}")
            #st.toast(f"Missing: {', '.join(readable)}", icon="⚠️")
        else:
            # All required fields are present – the user can start the workflow
            if st.button("Start", type="primary"):
                form_data = collect_form_data()
                st.session_state.form_data = form_data
                st.session_state.form_submitted = True
                st.session_state.unsaved_work = True

                # ---- Print to terminal only ----
                print("=== filled data sent to LLM ===")
                print(json.dumps(form_data, indent=2))
                #--Continue with LLm call--
                with st.spinner("Thinking…"):
                    first_q = call_ollama(build_question_prompt(form_data, [], {}))

                st.session_state.current_question = first_q
                st.rerun()

            # ── PHASE 2 : Chat ──
            elif st.session_state.form_submitted:
                for msg in st.session_state.conversation:
                    with st.chat_message(msg["role"]):
                        st.write(msg["content"])

                if st.session_state.justification_generated:
                    st.subheader("Generated Justification")
                    st.write(st.session_state.final_justification)

                    buf = generate_docx(
                        st.session_state.form_data.get("item_nomenclature", ""),
                        st.session_state.final_justification,
                    )
                    fname = (
                        f"{st.session_state.form_data.get('item_nomenclature','item').replace(' ', '_')}.docx"
                    )
                    st.download_button(
                        "Download DOCX",
                        data=buf,
                        file_name=fname,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )

                else:
                    with st.chat_message("assistant"):
                        st.write(st.session_state.current_question)

                    user_answer = st.chat_input("Enter your answer")

                    if user_answer:
                        st.session_state.conversation.append(
                            {"role": "assistant", "content": st.session_state.current_question}
                        )
                        st.session_state.conversation.append(
                            {"role": "user", "content": user_answer}
                        )
                        qkey = f"question_{st.session_state.question_count + 1}"
                        st.session_state.answers[qkey] = {
                            "question": st.session_state.current_question,
                            "answer": user_answer,
                        }
                        st.session_state.question_count += 1
                        st.session_state.unsaved_work = True

                        with st.spinner("Thinking..."):
                            next_q = call_ollama(build_question_prompt(
                                st.session_state.form_data,
                                st.session_state.conversation,
                                st.session_state.answers,
                            ))

                        if "JUSTIFICATION_COMPLETE" in next_q:
                            with st.spinner("Generating justification..."):
                                final = call_ollama(build_justification_prompt(
                                    st.session_state.form_data,
                                    st.session_state.answers,
                                ))
                            st.session_state.final_justification = final
                            st.session_state.justification_generated = True
                            st.session_state.unsaved_work = True
                        else:
                            st.session_state.current_question = next_q

                        st.rerun()


# ══════════════════════════════════════════════════════════
#  ENTRY POINT
# ══════════════════════════════════════════════════════════

# Restore login from cookie
if not st.session_state.logged_in and COOKIES_OK:
    _u = _get_cookie("username")
    _i = _get_cookie("user_id")
    if _u and _i:
        st.session_state.logged_in = True
        st.session_state.username  = _u
        st.session_state.user_id = int(_i) if _i.isdigit() else None
        refresh_sessions()

if not st.session_state.logged_in:
    show_login()

else:
    show_sidebar()
    show_main_app()

    # Modal dialogs — rendered as overlays on top of the page
    if st.session_state.show_save_dialog:
        save_dialog()

    if st.session_state.rename_id is not None:
        rename_dialog()

    if st.session_state.delete_id is not None:
        delete_dialog()

    if st.session_state.pending_new_session:
        unsaved_dialog()
